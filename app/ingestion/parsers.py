from __future__ import annotations

from datetime import datetime
from typing import Protocol

import fitz
from docx import Document as DocxDocument
from markdown_it import MarkdownIt

from app.documents.schemas import ParsedDocument, ParsedSection


class DocumentParser(Protocol):
    def parse(
        self, document_id: str, tenant_id: str, source_path: str, title: str | None = None
    ) -> ParsedDocument: ...


class PDFParser:
    def parse(
        self, document_id: str, tenant_id: str, source_path: str, title: str | None = None
    ) -> ParsedDocument:
        pdf = fitz.open(source_path)
        sections: list[ParsedSection] = []
        extracted_pages: list[str] = []
        for index, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if not text:
                continue
            extracted_pages.append(text)
            sections.append(
                ParsedSection(
                    section_id=f"{document_id}:p{index}",
                    content=text,
                    page_number=index,
                    heading=f"Page {index}",
                    metadata={"page_number": index},
                )
            )
        return ParsedDocument(
            document_id=document_id,
            tenant_id=tenant_id,
            title=title,
            source_type="pdf",
            source_path=source_path,
            metadata={"page_count": len(sections)},
            sections=sections,
            extracted_text="\n\n".join(extracted_pages),
            created_at=datetime.utcnow(),
        )


class MarkdownParser:
    def __init__(self) -> None:
        self.md = MarkdownIt()

    def parse(
        self, document_id: str, tenant_id: str, source_path: str, title: str | None = None
    ) -> ParsedDocument:
        with open(source_path, "r", encoding="utf-8") as f:
            text = f.read()
        tokens = self.md.parse(text)
        sections: list[ParsedSection] = []
        current_heading: str | None = None
        buffer: list[str] = []
        counter = 0
        for token in tokens:
            if token.type == "heading_open":
                if buffer:
                    counter += 1
                    sections.append(
                        ParsedSection(
                            section_id=f"{document_id}:s{counter}",
                            content="\n".join(buffer).strip(),
                            heading=current_heading,
                            metadata={"heading": current_heading},
                        )
                    )
                    buffer = []
            elif token.type == "inline":
                content = token.content.strip()
                if content:
                    if current_heading is None:
                        current_heading = content if len(content) < 120 else None
                    else:
                        buffer.append(content)
        if buffer:
            counter += 1
            sections.append(
                ParsedSection(
                    section_id=f"{document_id}:s{counter}",
                    content="\n".join(buffer).strip(),
                    heading=current_heading,
                    metadata={"heading": current_heading},
                )
            )
        if not sections:
            sections = [ParsedSection(section_id=f"{document_id}:s1", content=text, metadata={})]
        return ParsedDocument(
            document_id=document_id,
            tenant_id=tenant_id,
            title=title,
            source_type="md",
            source_path=source_path,
            metadata={"section_count": len(sections)},
            sections=sections,
            extracted_text=text,
            created_at=datetime.utcnow(),
        )


class TXTParser:
    def parse(
        self, document_id: str, tenant_id: str, source_path: str, title: str | None = None
    ) -> ParsedDocument:
        with open(source_path, "r", encoding="utf-8") as f:
            text = f.read()
        return ParsedDocument(
            document_id=document_id,
            tenant_id=tenant_id,
            title=title,
            source_type="txt",
            source_path=source_path,
            metadata={},
            sections=[ParsedSection(section_id=f"{document_id}:s1", content=text, metadata={})],
            extracted_text=text,
            created_at=datetime.utcnow(),
        )


class DOCXParser:
    def parse(
        self, document_id: str, tenant_id: str, source_path: str, title: str | None = None
    ) -> ParsedDocument:
        doc = DocxDocument(source_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(
            document_id=document_id,
            tenant_id=tenant_id,
            title=title,
            source_type="docx",
            source_path=source_path,
            metadata={"paragraph_count": len(paragraphs)},
            sections=[ParsedSection(section_id=f"{document_id}:s1", content=text, metadata={})],
            extracted_text=text,
            created_at=datetime.utcnow(),
        )
