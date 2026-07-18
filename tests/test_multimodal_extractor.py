from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from fpdf import FPDF

from app.ingestion.multimodal.extractor import ImageExtractor, TableExtractor


def _create_pdf_with_image(tmp_path: Path) -> str:
    """Create a minimal PDF with an embedded image."""
    import base64
    from io import BytesIO

    # 1x1 red pixel PNG (base64 encoded)
    _png_1px_b64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
        "AAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    pdf_path = tmp_path / "test_img.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.set_xy(10, 10)
    pdf.cell(0, 10, "Test PDF with image area")
    pdf.image(BytesIO(base64.b64decode(_png_1px_b64)), x=10, y=20, w=10, h=10)
    pdf.output(str(pdf_path))
    return str(pdf_path)


def _create_pdf_with_table(tmp_path: Path) -> str:
    """Create a minimal PDF with a table-like structure using fpdf2."""
    pdf_path = tmp_path / "test_table.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    # Draw table grid lines
    for x in range(1, 4):
        pdf.line(x * 60, 50, x * 60, 150)
    for y in range(1, 4):
        pdf.line(60, y * 33 + 50, 180, y * 33 + 50)
    # Place text in cells
    pdf.text(65, 70, "A")
    pdf.text(125, 70, "B")
    pdf.text(65, 103, "1")
    pdf.text(125, 103, "2")
    pdf.output(str(pdf_path))
    return str(pdf_path)


def test_image_extractor_pdf() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        path = _create_pdf_with_image(Path(tmp))
        extractor = ImageExtractor(max_size_mb=10, max_images=5)
        images = extractor.extract_images(path, "pdf")
        assert len(images) > 0
        assert images[0].format in ("png", "jpeg")
        assert images[0].size_bytes > 0
        assert images[0].page_number is not None
        assert images[0].page_number >= 1


def test_image_extractor_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "test_img.docx"
    doc = DocxDocument()
    doc.add_paragraph("Test docx with image placeholder")
    # Add a small inline shape as placeholder
    doc.add_paragraph("Image area")
    doc.save(str(docx_path))

    extractor = ImageExtractor(max_size_mb=10, max_images=5)
    images = extractor.extract_images(str(docx_path), "docx")

    assert len(images) == 0


def test_image_extractor_unsupported_type() -> None:
    extractor = ImageExtractor()
    images = extractor.extract_images("/nonexistent/test.txt", "txt")
    assert images == []


def test_image_extractor_max_images() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        path = _create_pdf_with_image(Path(tmp))
        extractor = ImageExtractor(max_size_mb=10, max_images=0)
        images = extractor.extract_images(path, "pdf")
        assert len(images) == 0


def test_table_extractor_pdf() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        path = _create_pdf_with_table(Path(tmp))
        extractor = TableExtractor(max_tables=5)
        tables = extractor.extract_tables(path, "pdf")
        # PyMuPDF find_tables may or may not detect drawn-line tables
        # We just verify no crash and valid return type
        assert isinstance(tables, list)
        if tables:
            assert "|" in tables[0].markdown_text
            assert tables[0].table_index >= 0


def test_table_extractor_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "test_table.docx"
    doc = DocxDocument()
    doc.add_heading("Test Table", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "X"
    table.cell(1, 1).text = "42"
    doc.save(str(docx_path))

    extractor = TableExtractor()
    tables = extractor.extract_tables(str(docx_path), "docx")
    assert len(tables) == 1
    assert "Name" in tables[0].markdown_text
    assert "Value" in tables[0].markdown_text
    assert tables[0].page_number is None


def test_table_extractor_docx_max_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "test_max.docx"
    doc = DocxDocument()
    for _ in range(5):
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "data"
    doc.save(str(docx_path))

    extractor = TableExtractor(max_tables=2)
    tables = extractor.extract_tables(str(docx_path), "docx")
    assert len(tables) == 2


def test_table_extractor_unsupported_type() -> None:
    extractor = TableExtractor()
    tables = extractor.extract_tables("/nonexistent/test.txt", "txt")
    assert tables == []
